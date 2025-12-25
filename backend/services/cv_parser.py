"""
Universal CV Parser Service
Extracts text and skills from any domain/industry CV
"""
import fitz  # PyMuPDF
from docx import Document
import re
from typing import Dict, List, Any, Optional
import io
from datetime import datetime


import pytesseract
from pdf2image import convert_from_bytes

class CVParser:
    """Service for parsing and extracting information from CVs (all industries)"""
    
    # EXPANDED: Skills from ALL domains
    UNIVERSAL_SKILLS = {
        # === TECHNICAL/IT ===
        # Programming
        "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
        "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "sql",
        
        # AI/ML
        "machine learning", "deep learning", "neural networks", "nlp",
        "computer vision", "tensorflow", "pytorch", "keras", "scikit-learn",
        "transformers", "llm", "rag", "langchain",
        
        # Web Dev
        "react", "vue", "angular", "node.js", "express", "fastapi", "django",
        "flask", "spring", "html", "css", "tailwind", "bootstrap",
        
        # Cloud & DevOps
        "aws", "azure", "gcp", "docker", "kubernetes", "terraform",
        "jenkins", "ci/cd", "git", "github", "gitlab",
        
        # Databases
        "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
        "dynamodb", "firebase", "sqlite", "oracle",
        
        # === HEALTHCARE ===
        "patient care", "clinical documentation", "hipaa", "emr", "ehr",
        "medical terminology", "vital signs", "cpr", "bls", "acls",
        "patient assessment", "medication administration", "iv therapy",
        "wound care", "electronic health records", "medical coding",
        "icd-10", "healthcare compliance",
        
        # === MARKETING ===
        "seo", "google analytics", "content strategy", "social media marketing",
        "email marketing", "ppc", "google ads", "facebook ads", "hubspot",
        "marketing automation", "conversion optimization", "a/b testing",
        "content creation", "copywriting", "brand management", "crm",
        "salesforce", "market research",
        
        # === FINANCE/ACCOUNTING ===
        "financial modeling", "gaap", "ifrs", "financial analysis",
        "budgeting", "forecasting", "excel", "quickbooks", "sap",
        "accounts payable", "accounts receivable", "audit", "tax preparation",
        "risk assessment", "bloomberg terminal", "financial reporting",
        "variance analysis", "cost accounting", "cpa", "cfa",
        
        # === SALES ===
        "lead generation", "cold calling", "relationship building",
        "negotiation", "sales pipeline", "quota achievement", "crm software",
        "b2b sales", "b2c sales", "account management", "upselling",
        "customer retention", "sales presentations",
        
        # === HUMAN RESOURCES ===
        "recruitment", "talent acquisition", "onboarding", "employee relations",
        "performance management", "hris", "workday", "adp", "payroll",
        "benefits administration", "training and development", "hr compliance",
        "labor law", "employee engagement", "compensation analysis",
        
        # === PROJECT MANAGEMENT ===
        "agile", "scrum", "kanban", "waterfall", "project planning",
        "risk management", "stakeholder management", "budget management",
        "resource allocation", "jira", "confluence", "ms project",
        "pmp", "prince2", "gantt charts",
        
        # === DESIGN ===
        "adobe photoshop", "adobe illustrator", "figma", "sketch",
        "ui design", "ux design", "wireframing", "prototyping",
        "user research", "visual design", "graphic design", "branding",
        "typography", "color theory", "adobe xd",
        
        # === EDUCATION ===
        "curriculum development", "lesson planning", "classroom management",
        "student assessment", "differentiated instruction", "educational technology",
        "learning management systems", "google classroom", "canvas",
        "special education", "tesol", "esl", "teaching certification",
        
        # === LEGAL ===
        "contract law", "litigation", "legal research", "legal writing",
        "case management", "westlaw", "lexisnexis", "compliance",
        "corporate law", "intellectual property", "employment law",
        "regulatory compliance", "due diligence",
        
        # === MANUFACTURING/ENGINEERING ===
        "autocad", "solidworks", "cad", "lean manufacturing", "six sigma",
        "quality control", "iso 9001", "osha", "process improvement",
        "supply chain", "inventory management", "plc programming",
        "cnc", "welding", "blueprint reading",
        
        # === SOFT SKILLS (Universal) ===
        "leadership", "communication", "problem solving", "teamwork",
        "time management", "critical thinking", "adaptability",
        "conflict resolution", "presentation skills", "analytical skills",
        "attention to detail", "customer service", "multitasking",
        "decision making", "collaboration", "interpersonal skills"
    }
    
    # Keep education keywords
    EDUCATION_KEYWORDS = [
        "bachelor", "master", "phd", "doctorate", "degree", "university",
        "college", "bsc", "msc", "ba", "ma", "mba", "engineering",
        "diploma", "certification", "associate", "graduate",
        "a level", "a-level", "gcse", "school", "sixth form", "academy"
    ]
    
    def parse_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """Parse PDF with OCR fallback for image-only files"""
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            page_count = len(doc)
            doc.close()
            
            # OCR Fallback: If text is minimal (<100 chars), try OCR
            if len(text.strip()) < 100:
                print("Low text density detected in PDF. Attempting OCR...")
                try:
                    images = convert_from_bytes(file_content)
                    ocr_text = ""
                    for i, image in enumerate(images):
                        # Preprocessing: Convert to grayscale
                        gray_image = image.convert('L')
                        ocr_text += pytesseract.image_to_string(gray_image) + "\n"
                    
                    if len(ocr_text.strip()) > len(text.strip()):
                        text = ocr_text
                        print(f"OCR successful. Extracted {len(text)} characters.")
                except Exception as ocr_e:
                    print(f"OCR failed: {str(ocr_e)}")
                    # Continue with whatever text we found originally
            
            return {
                "success": True,
                "text": text,
                "page_count": page_count,
                "format": "pdf"
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "format": "pdf"
            }
    
    def parse_docx(self, file_content: bytes) -> Dict[str, Any]:
        """Parse DOCX - unchanged"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return {
                "success": True,
                "text": text,
                "paragraph_count": len(doc.paragraphs),
                "format": "docx"
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "format": "docx"
            }
    
    def parse_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse any file - unchanged"""
        filename_lower = filename.lower()
        if filename_lower.endswith(".pdf"):
            return self.parse_pdf(file_content)
        elif filename_lower.endswith(".docx"):
            return self.parse_docx(file_content)
        elif filename_lower.endswith(".txt"):
            try:
                text = file_content.decode("utf-8")
                return {"success": True, "text": text, "format": "txt"}
            except Exception as e:
                return {"success": False, "error": str(e), "text": "", "format": "txt"}
        else:
            return {
                "success": False,
                "error": f"Unsupported file format: {filename}",
                "text": "",
                "format": "unknown"
            }
    
    def extract_skills(self, text: str) -> List[str]:
        """
        Extract skills from ANY domain CV
        Now works for technical, healthcare, marketing, finance, etc.
        """
        text_lower = text.lower()
        found_skills = []
        
        for skill in self.UNIVERSAL_SKILLS:
            if " " not in skill:
                # Single-word skill: use word boundary
                pattern = rf'\b{re.escape(skill)}\b'
                if re.search(pattern, text_lower):
                    found_skills.append(skill.title())
            else:
                # Multi-word skill: substring match
                if skill in text_lower:
                    found_skills.append(skill.title())
        
        return sorted(set(found_skills))
    
    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact info - unchanged"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        
        # Improved phone regex: supports +44-XXXX-XXXXXX, +1 (234) 567-8901, etc.
        phone_pattern = r'(?:\+\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?)?\d{3,4}[\s.-]?\d{3,4}[\s.-]?\d{0,4}'
        phone_match = re.search(phone_pattern, text)
        
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text.lower())
        
        return {
            "email": email_match.group() if email_match else None,
            "phone": phone_match.group() if phone_match else None,
            "linkedin": linkedin_match.group() if linkedin_match else None
        }
    
    def extract_education(self, text: str) -> List[str]:
        """Extract education - unchanged"""
        lines = text.split("\n")
        education_lines = []
        
        for line in lines:
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in self.EDUCATION_KEYWORDS):
                if len(line.strip()) > 10:
                    education_lines.append(line.strip())
        
        return education_lines[:5]
    
    def extract_name(self, text: str) -> Optional[str]:
        """
        Extract candidate name from CV (usually in first 3 lines)
        """
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        
        for line in lines[:3]:
            # Skip lines with email or phone numbers
            if "@" in line or re.search(r'\d{3}[-.\s]?\d{3}', line):
                continue
            
            words = line.split()
            # Name is typically 2-4 words, less than 50 chars
            if 2 <= len(words) <= 4 and len(line) < 50:
                # Check if mostly alphabetic characters
                alpha_ratio = sum(c.isalpha() or c.isspace() for c in line) / len(line)
                if alpha_ratio > 0.8:
                    return line
        
        return None
    
    def extract_experience_years(self, text: str) -> Optional[float]:
        """
        Calculate total years of work experience from date ranges
        Looks for patterns like: "2020 - 2023" or "Jan 2020 - Present"
        """
        # Pattern matches: 2020 - 2023, 2020 - Present, etc.
        date_pattern = r'(\d{4})\s*[-–—]\s*(?:(\d{4})|present|current)'
        matches = re.findall(date_pattern, text.lower())
        
        if not matches:
            return None
        
        current_year = datetime.now().year
        total_years = 0
        
        for match in matches:
            start_year = int(match[0])
            end_year = int(match[1]) if match[1] else current_year
            years = end_year - start_year
            
            # Sanity check: experience should be 0-50 years
            if 0 <= years <= 50:
                total_years += years
        
        return round(total_years, 1) if total_years > 0 else None
    
    def detect_industry(self, skills: List[str], text: str) -> str:
        """
        Detect primary industry/domain based on skills and CV content
        """
        text_lower = text.lower()
        
        industry_keywords = {
            "Software Engineering": ["python", "javascript", "react", "api", "git", "docker", "programming"],
            "Data Science/AI": ["machine learning", "tensorflow", "data analysis", "statistics", "data science"],
            "Healthcare": ["patient care", "clinical", "medical", "hipaa", "emr", "nursing", "healthcare"],
            "Marketing": ["seo", "marketing", "content", "social media", "google analytics", "campaign"],
            "Finance": ["financial", "accounting", "gaap", "audit", "excel", "budgeting", "finance"],
            "Design": ["photoshop", "figma", "ui", "ux", "design", "visual", "graphic"],
            "HR": ["recruitment", "hr", "hiring", "onboarding", "employee", "human resources"],
            "Sales": ["sales", "crm", "lead generation", "b2b", "negotiation", "revenue"],
            "Education": ["teaching", "curriculum", "classroom", "student", "education", "instructor"]
        }
        
        scores = {}
        for industry, keywords in industry_keywords.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            scores[industry] = score
        
        # Return industry with highest score, or "General" if score too low
        detected = max(scores, key=scores.get)
        return detected if scores[detected] >= 2 else "General"
    
    def analyze_cv(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Fully analyze a CV file
        NOW INCLUDES: name, experience years, and industry detection
        """
        parse_result = self.parse_file(file_content, filename)
        
        if not parse_result["success"]:
            return {
                "success": False,
                "error": parse_result.get("error", "Failed to parse file")
            }
        
        text = parse_result["text"]
        skills = self.extract_skills(text)
        contact = self.extract_contact_info(text)
        education = self.extract_education(text)
        name = self.extract_name(text)
        experience_years = self.extract_experience_years(text)
        industry = self.detect_industry(skills, text)
        
        return {
            "success": True,
            "filename": filename,
            "format": parse_result["format"],
            "text": text,
            "text_length": len(text),
            "name": name,
            "skills": skills,
            "skills_count": len(skills),
            "contact": contact,
            "education": education,
            "experience_years": experience_years,
            "detected_industry": industry
        }


# Singleton
cv_parser = CVParser()
